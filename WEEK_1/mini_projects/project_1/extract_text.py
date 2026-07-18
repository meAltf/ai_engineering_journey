from PyPDF2 import PdfReader
from docx import Document
import json

# Read PDF file and extract text
def read_pdf(file_path):
    pdf_reader = PdfReader(file_path)
    text = ""

    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"

    return text


# Read Word document and extract text
def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text

# Handles both formats dynamically based on the file extension
def read_resume(file_path):
    if file_path.endswith(".pdf"):
        return read_pdf(file_path)
    elif file_path.endswith(".docx"):
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a PDF or Word document.")

# just to resolve circular dependency between resume_extract & prompts files
def get_resume_text(file_path):
    return read_resume(file_path)



#pdf_text = read_resume("resumes/fullstack_resume_pdf.pdf")
# print("PDF resume:", pdf_text[:10000])